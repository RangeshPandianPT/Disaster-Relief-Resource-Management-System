import docx
from docx.shared import Pt, Inches
import os

def generate_report(template_path, output_path):
    print(f"Loading template from: {template_path}")
    doc = docx.Document(template_path)

    # Content to insert
    content_map = {
        "Introduction": (
            "The Disaster Relief Resource Management System (DRRMS) is a comprehensive database management system designed to coordinate disaster relief operations efficiently. "
            "It manages critical data regarding disasters, affected areas, resources, inventory, relief teams, volunteers, and donations."
        ),
        "Motivation": (
            "In the wake of natural or man-made disasters, rapid response is crucial to saving lives and minimizing damage. "
            "Manual coordination of resources and teams is often slow, error-prone, and lacks real-time visibility. "
            "There is a critical need for a centralized system to streamline these operations, ensuring that the right resources reach the right people at the right time. "
            "This project aims to solve these logistical challenges through a robust database solution."
        ),
        "Identification of Entity and Relationships": (
            "The system is built upon the following core entities:\n"
            "- DISASTER: The central event triggering system activity.\n"
            "- AFFECTED_AREA: Specific locations impacted by a disaster.\n"
            "- RESOURCE: Types of aid items available.\n"
            "- INVENTORY: Actual stock of resources in warehouses.\n"
            "- RELIEF_TEAM: Groups deployed to help.\n"
            "- VOLUNTEER: Individuals helping out (linked to teams).\n"
            "- REQUEST: Demands for resources from affected areas.\n"
            "- ALLOCATION: Fulfillment of requests from inventory.\n"
            "- DONOR and DONATION: Tracking incoming aid.\n\n"
            "Key Relationships:\n"
            "- One Disaster affects multiple Areas (1:N).\n"
            "- Resources are stored in Inventory and allocated to Requests (N:M via allocation).\n"
            "- Relief Teams consist of multiple Volunteers (1:N)."
        ),
        "Scope": (
            "The scope of the DRRMS includes:\n"
            "- Disaster Management: Tracking details of various disaster events and their severity.\n"
            "- Resource Tracking: Managing inventory of essential supplies (food, medicine, equipment) across multiple warehouses.\n"
            "- Relief Operations: Coordinating relief teams and assigning them to affected areas.\n"
            "- Volunteer & Donation Management: tracking volunteer skills/availability and donor contributions.\n"
            "- Reporting: Generating insights on resource utilization and delivering status updates."
        ),
        "Problem Statement": (
            "Current disaster relief efforts often suffer from:\n"
            "1. Lack of real-time inventory visibility, leading to shortages or wastage.\n"
            "2. Inefficient deployment of relief teams due to poor communication.\n"
            "3. Difficulty in matching volunteer skills with immediate needs.\n"
            "4. Fragmented data storage making it hard to generate comprehensive situation reports.\n"
            "This project addresses these issues by providing a unified platform for data management."
        ),
        "Project Requirements": (
            "The system requires:\n"
            "1. Database Backend: A normalized relational database (MySQL) to store complex relationships between entities.\n"
            "2. Command Line Interface (CLI): For administrators to manage the system efficiently with low resource overhead.\n"
            "3. Web Application: A user-friendly dashboard for real-time monitoring and visualization.\n"
            "4. Data Integrity: Enforced via foreign keys, triggers, and transactions.\n"
            "5. Security: Role-based access control to protect sensitive data."
        ),
        "Construction of DB Using ER Model": (
            "The Entity-Relationship (ER) diagram models the logical structure of the database. "
            "It defines the entities, their attributes, and the relationships between them. "
            "Please refer to the 'Identification of Entity and Relationships' section for a detailed description of the components shown in the diagram above."
        ),
        "Design of Relational Schemas": (
            "The relational schema for the DRRMS project is depicted below. It outlines the tables, columns, and foreign key constraints corresponding to the entities defined previously."
        )
    }

    # Iterate through paragraphs and insert content after headers
    # Note: We collect the paragraphs first because inserting modifies the list structure?
    # python-docx paragraphs list is live. But inserting via `insert_paragraph_before` might shift indices.
    # It's safer to not rely on index stability if we are inserting.
    # However, our previous approach using `range(len(paragraphs))` and checking `text` works if we break or handle shifts.
    # But since we only insert *once* per header and they are distinct, simple iteration might miss things if we insert before the *next* one.
    
    # We will assume headers are unique and distinct.
    
    doc_paragraphs = doc.paragraphs # Reference to live list
    
    # We iterate backwards to avoid index shifting issues? No, we are searching for text.
    # Let's stick to the previous loop which worked, just modify the inside.
    
    i = 0
    while i < len(doc_paragraphs):
        p = doc_paragraphs[i]
        text = p.text.strip()
        
        matched_key = None
        if "Introduction" == text: matched_key = "Introduction"
        elif "Motivation" == text: matched_key = "Motivation"
        elif "Scope" == text: matched_key = "Scope"
        elif "Problem Statement" == text: matched_key = "Problem Statement"
        elif "Project Requirements" == text: matched_key = "Project Requirements"
        elif "Identification of Entity and Relationships" in text: matched_key = "Identification of Entity and Relationships"
        elif "Construction of DB Using ER Model" in text: matched_key = "Construction of DB Using ER Model"
        elif "Design of Relational Schemas" in text: matched_key = "Design of Relational Schemas"

        if matched_key:
            print(f"Populating section: {matched_key}")
            # Target is the pixel after the header. 
            # If we are at 'i', we want to insert after 'i'.
            # So we insert before 'i+1'.
            
            if i + 1 < len(doc_paragraphs):
                target_p = doc_paragraphs[i+1]
                
                # Special handling for ER Model Image or Relational Schema Image
                if matched_key == "Construction of DB Using ER Model" or matched_key == "Design of Relational Schemas":
                    image_path = r"docs\er_diagram.png"
                    if os.path.exists(image_path):
                        print(f"Inserting image: {image_path}")
                        img_p = target_p.insert_paragraph_before()
                        run = img_p.add_run()
                        run.add_picture(image_path, width=Inches(6))
                        
                        txt_p = target_p.insert_paragraph_before(content_map[matched_key])
                        txt_p.style = 'Normal'
                    else:
                        print(f"Warning: Image {image_path} not found.")
                        new_p = target_p.insert_paragraph_before(content_map[matched_key])
                        new_p.style = 'Normal'
                else:
                    new_p = target_p.insert_paragraph_before(content_map[matched_key])
                    new_p.style = 'Normal'
            else:
                # End of doc
                if matched_key == "Construction of DB Using ER Model" or matched_key == "Design of Relational Schemas":
                     image_path = r"docs\er_diagram.png"
                     if os.path.exists(image_path):
                        doc.add_picture(image_path, width=Inches(6))
                doc.add_paragraph(content_map[matched_key])
        
        i += 1

    doc.save(output_path)
    print(f"Report saved to {output_path}")

if __name__ == "__main__":
    generate_report("Report template.docx", "DRRMS_Report_v2.docx")


