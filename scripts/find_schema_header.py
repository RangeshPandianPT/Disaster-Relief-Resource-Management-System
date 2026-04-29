import docx

def find_relational_schema(file_path):
    try:
        doc = docx.Document(file_path)
        print(f"Searching for 'Relational Schema' in {file_path}...")
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if "Relational Schema" in text or "2.1" in text:
                print(f"Found potential header at P{i}: {text[:100]}")
                # Print context
                for j in range(1, 4):
                    if i+j < len(doc.paragraphs):
                         print(f"  Context P{i+j}: {doc.paragraphs[i+j].text[:50]}...")

    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    find_relational_schema(r"Report template.docx")
