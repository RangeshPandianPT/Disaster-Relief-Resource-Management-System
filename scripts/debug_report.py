import docx

def debug_report(file_path):
    print(f"Debugging {file_path}...")
    try:
        doc = docx.Document(file_path)
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if "Identification of Entity and Relationships" in text:
                print(f"--- Found Header at P{i} ---")
                print(f"P{i}: {text}")
                # Print next 5 paragraphs
                for j in range(1, 6):
                    if i + j < len(doc.paragraphs):
                        print(f"P{i+j}: {doc.paragraphs[i+j].text[:100]}...")
                        
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    debug_report("DRRMS_Report.docx")
