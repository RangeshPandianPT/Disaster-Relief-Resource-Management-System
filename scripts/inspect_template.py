import docx

def read_docx_structure(file_path):
    try:
        doc = docx.Document(file_path)
        print(f"Structure of {file_path}:")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"Paragraph {i}: Style='{paragraph.style.name}', Text='{paragraph.text[:100]}...'")
        
        print("\nTables:")
        for i, table in enumerate(doc.tables):
            print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
            
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    read_docx_structure(r"d:\DBMS\DRRMS\Report template.docx")
